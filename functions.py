import telebot
import requests
import random
import datetime
import time
import os

from docx import Document

import db_functions
import config
import text


bot = telebot.TeleBot(config.TELEGRAM_TOKEN)


def get_good_answers_from_google():
    text = config.WORK_SHEET.col_values(1)[1::]
    recommended_products = config.WORK_SHEET.col_values(2)[1::]

    return list(zip(text, recommended_products))


def get_bad_answers_from_google():
    text = config.WORK_SHEET.col_values(3)[1::]

    return text


def get_average_answers_from_google():
    text = config.WORK_SHEET.col_values(6)[1::]

    return text


def get_greetings_from_google():
    text = config.WORK_SHEET.col_values(7)[1::]

    return text


def get_parting_from_google():
    text = config.WORK_SHEET.col_values(8)[1::]

    return text


def get_trigger_answers_from_google():
    text = config.WORK_SHEET.col_values(4)[1::]
    trigger = config.WORK_SHEET.col_values(5)[1::]

    return list(zip(text, trigger))


def select_unanswered():
    data = {
        'isAnswered' : False,
        'take' : 5000,
        'skip' : 0,
    }

    response = requests.get(url=config.API_ENDPOINT, params=data, headers=config.HEADERS_WB)
    
    if response.status_code == 200:
        try:
            feedbacks = response.json().get('data').get('feedbacks')
        except:
            feedbacks = None

        return feedbacks


def answer_unanswered():
    while True:
        if config.ACTIVE:
            db_functions.delete_answers('good')
            db_functions.delete_answers('bad')
            db_functions.delete_answers('trigger')

            good_answers = get_good_answers_from_google()
            for good_answer in good_answers:
                db_functions.add_good_answer(*good_answer)

            bad_answers = get_bad_answers_from_google()
            for bad_answer in bad_answers:
                db_functions.add_bad_answer(bad_answer)
            
            triggers = []
            trigger_answers = get_trigger_answers_from_google()
            for trigger_answer in trigger_answers:
                triggers.append(trigger_answer[1])
                db_functions.add_trigger_answer(*trigger_answer)

            average_answers = get_average_answers_from_google()
            greetings = get_greetings_from_google()
            partings = get_parting_from_google()

            feedbacks = select_unanswered()

            if feedbacks:
                for feedback in feedbacks:
                    review_id = feedback.get('id')
                    user_name = feedback.get('userName')
                    product_id = feedback.get('productDetails').get('nmId')
                    product_name = feedback.get('productDetails').get('productName')
                    review_mark = feedback.get('productValuation')
                    review_text = feedback.get('text')

                    for trigger in triggers:
                        if trigger in review_text.lower():
                            answer = db_functions.select_answer_by_trigger(trigger)
                            break
                    else:
                        if review_mark == 5:
                            answer = random.choice(db_functions.select_good_answers(product_id))
                        
                        elif review_mark == 4:
                            answer = random.choice(average_answers)
                            
                        else:
                            answer = random.choice(db_functions.select_bad_answers())

                    greeting = random.choice(greetings)
                    parting = random.choice(partings)

                    if user_name:
                        answer = f'{greeting}, {user_name.capitalize()}!\n' + answer + '\n\n' + parting
                    else:
                        answer = f'{greeting}!\n' + answer + '\n\n' + parting

                    answer_date = (datetime.datetime.utcnow() + datetime.timedelta(hours=3)).strftime("%d.%m.%Y")
                    
                    data = {
                        'id' : review_id,
                        'text' : answer,
                    }

                    try:
                        response = requests.patch(url=config.API_ENDPOINT, json=data, headers=config.HEADERS_WB)
                    except:
                        response = None

                    if response and response.status_code:
                        db_functions.add_answer_to_report(review_id, product_id, product_name, review_mark, review_text, answer, answer_date)
        
        time.sleep(1800)


def send_daily_report():
    while True:
        if (datetime.datetime.utcnow() + datetime.timedelta(hours=3)).strftime("%H:%M") == '23:58':
            current_date = (datetime.datetime.utcnow() + datetime.timedelta(hours=3)).strftime("%d.%m.%Y")
            date_name = (datetime.datetime.utcnow() + datetime.timedelta(hours=3)).strftime("%d-%m-%Y")

            answers_info = db_functions.select_today_answers_to_report(current_date)

            if answers_info:
                doc = Document()

                for answer_info in answers_info:
                    review_text = answer_info[5]
                    if not review_text:
                        review_text = 'без текста'

                    if answer_info[4] == 5:
                        doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
                
                for answer_info in answers_info:
                    review_text = answer_info[5]
                    if not review_text:
                        review_text = 'без текста'

                    if answer_info[4] == 4:
                        doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
                
                for answer_info in answers_info:
                    review_text = answer_info[5]
                    if not review_text:
                        review_text = 'без текста'

                    if answer_info[4] == 3:
                        doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
                
                for answer_info in answers_info:
                    review_text = answer_info[5]
                    if not review_text:
                        review_text = 'без текста'

                    if answer_info[4] == 2:
                        doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
                
                for answer_info in answers_info:
                    review_text = answer_info[5]
                    if not review_text:
                        review_text = 'без текста'

                    if answer_info[4] == 1:
                        doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')

                doc.save(f'{date_name}.docx')

                with open(f'{date_name}.docx', "rb") as f:
                    file_data = f.read()

                bot.send_document(chat_id=config.MANAGER_ID,
                                document=file_data,
                                visible_file_name=f'{date_name}.docx',
                                disable_notification=True, 
                                )

                try:
                    os.remove(f'{date_name}.docx')
                except:
                    pass

            else:
                bot.send_message(chat_id=config.MANAGER_ID,
                         text=text.no_data(current_date),
                         disable_notification=True,
                         )
            
            time.sleep(86350)


def send_date_report(current_date):
    date_name = current_date.replace('.', '-')

    answers_info = db_functions.select_today_answers_to_report(current_date)

    if answers_info:
        doc = Document()

        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 5:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 4:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 3:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 2:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 1:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        doc.save(f'{date_name}.docx')

        with open(f'{date_name}.docx', "rb") as f:
            file_data = f.read()

        bot.send_document(chat_id=config.MANAGER_ID,
                        document=file_data,
                        visible_file_name=f'{date_name}.docx', 
                        )
        
        try:
            os.remove(f'{date_name}.docx')
        except:
            pass

    else:
        bot.send_message(chat_id=config.MANAGER_ID,
                    text=text.no_data(current_date),
                    )


def send_all_report():
    answers_info = db_functions.select_all_answers_to_report()

    if answers_info:
        doc = Document()

        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 5:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 4:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 3:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 2:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        for answer_info in answers_info:
            review_text = answer_info[5]
            if not review_text:
                review_text = 'без текста'

            if answer_info[4] == 1:
                doc.add_paragraph(f'Дата ответа: {answer_info[7]}\nТовар: {answer_info[3]} ({answer_info[2]})\nОценка: {answer_info[4]}\nОтзыв: {review_text}\nОтвет: {answer_info[6]}\n\n')
        
        doc.save('all_answers.docx')

        with open('all_answers.docx', "rb") as f:
            file_data = f.read()

        bot.send_document(chat_id=config.MANAGER_ID,
                        document=file_data,
                        visible_file_name=f'all_answers.docx', 
                        )
        
        try:
            os.remove('all_answers.docx')
        except:
            pass

    else:
        bot.send_message(chat_id=config.MANAGER_ID,
                    text=text.NO_ANSWERS,
                    )